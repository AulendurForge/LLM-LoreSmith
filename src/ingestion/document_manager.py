"""
Document ingestion module for LLM LoreSmith.
Handles document upload, validation, and secure storage.
"""

import os
import hashlib
import shutil
import filetype
import PyPDF2
from docx import Document
from typing import Dict, List, Optional, Tuple, BinaryIO
from pathlib import Path
from datetime import datetime

from ..config.config_manager import config


class DocumentValidator:
    """Validates documents for credibility and quality."""
    
    # Supported document types and their MIME types
    SUPPORTED_TYPES = {
        'application/pdf': '.pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'text/plain': '.txt',
        'text/markdown': '.md',
        'text/html': '.html',
    }
    
    # Minimum document size in bytes (1KB)
    MIN_SIZE = 1024
    
    # Maximum document size in bytes (from config, default 50MB)
    @property
    def max_size(self) -> int:
        """Get maximum document size in bytes."""
        return config.get('resources.max_document_size_mb', 50) * 1024 * 1024
    
    def validate_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate a document for type, size, and content quality.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (is_valid, message)
        """
        # Check if file exists
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size < self.MIN_SIZE:
            return False, f"File is too small (minimum {self.MIN_SIZE} bytes)"
        if file_size > self.max_size:
            return False, f"File is too large (maximum {self.max_size} bytes)"
        
        # Check file type using filetype instead of magic
        kind = filetype.guess(file_path)
        if kind is None:
            # For text files that might not be detected by filetype
            # Try to determine by extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.txt', '.md', '.html']:
                mime_type = {
                    '.txt': 'text/plain',
                    '.md': 'text/markdown', 
                    '.html': 'text/html'
                }.get(ext)
            else:
                return False, "Unknown file type"
        else:
            mime_type = kind.mime
            
        if mime_type not in self.SUPPORTED_TYPES:
            return False, f"Unsupported file type: {mime_type}"
        
        # Check content quality based on file type
        if mime_type == 'application/pdf':
            return self._validate_pdf(file_path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._validate_docx(file_path)
        elif mime_type in ['text/plain', 'text/markdown', 'text/html']:
            return self._validate_text(file_path)
        
        return True, "Document is valid"
    
    def _validate_pdf(self, file_path: str) -> Tuple[bool, str]:
        """Validate PDF document content."""
        try:
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                if len(pdf.pages) < 1:
                    return False, "PDF document has no pages"
                
                # Check if PDF has text content
                text_content = ""
                for i in range(min(3, len(pdf.pages))):
                    text_content += pdf.pages[i].extract_text()
                
                if not text_content.strip():
                    return False, "PDF document appears to have no text content"
                
                return True, "PDF document is valid"
        except Exception as e:
            return False, f"Error validating PDF: {str(e)}"
    
    def _validate_docx(self, file_path: str) -> Tuple[bool, str]:
        """Validate DOCX document content."""
        try:
            doc = Document(file_path)
            if len(doc.paragraphs) < 1:
                return False, "DOCX document has no paragraphs"
            
            # Check if document has text content
            text_content = "\n".join([p.text for p in doc.paragraphs[:10]])
            if not text_content.strip():
                return False, "DOCX document appears to have no text content"
            
            return True, "DOCX document is valid"
        except Exception as e:
            return False, f"Error validating DOCX: {str(e)}"
    
    def _validate_text(self, file_path: str) -> Tuple[bool, str]:
        """Validate text document content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()[:20]  # Read first 20 lines
                
            if not lines:
                return False, "Text document is empty"
            
            text_content = "".join(lines)
            if not text_content.strip():
                return False, "Text document appears to have no content"
            
            return True, "Text document is valid"
        except Exception as e:
            return False, f"Error validating text document: {str(e)}"


class DocumentStorage:
    """Manages secure document storage."""
    
    def __init__(self, base_dir: Optional[str] = None):
        """
        Initialize document storage.
        
        Args:
            base_dir: Base directory for document storage. If None, uses default.
        """
        self.base_dir = base_dir or os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "data",
            "documents"
        )
        os.makedirs(self.base_dir, exist_ok=True)
    
    def store_document(self, file_path: str, metadata: Optional[Dict] = None) -> str:
        """
        Store a document securely.
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata for the document
            
        Returns:
            Document ID
        """
        # Generate document ID
        doc_id = self._generate_document_id(file_path)
        
        # Create document directory
        doc_dir = os.path.join(self.base_dir, doc_id)
        os.makedirs(doc_dir, exist_ok=True)
        
        # Copy document file
        ext = os.path.splitext(file_path)[1]
        if not ext:
            # Determine extension from file type if not in filename
            kind = filetype.guess(file_path)
            if kind is None:
                # For text files, try to determine by reading first few bytes
                with open(file_path, 'rb') as f:
                    data = f.read(4096)
                    if all(c < 128 for c in data):  # ASCII text
                        ext = '.txt'
                    else:
                        ext = '.bin'  # Binary file
            else:
                ext = self.document_validator.SUPPORTED_TYPES.get(kind.mime, '.bin')
                
        dest_path = os.path.join(doc_dir, f"document{ext}")
        shutil.copy2(file_path, dest_path)
        
        # Store metadata
        if metadata is None:
            metadata = {}
        
        # Add system metadata
        metadata.update({
            "document_id": doc_id,
            "filename": os.path.basename(file_path),
            "size_bytes": os.path.getsize(file_path),
            "storage_time": datetime.now().isoformat(),
        })
        
        # Save metadata
        self._save_metadata(doc_id, metadata)
        
        return doc_id
    
    def get_document_path(self, doc_id: str) -> Optional[str]:
        """
        Get the path to a stored document.
        
        Args:
            doc_id: Document ID
            
        Returns:
            Path to the document file or None if not found
        """
        doc_dir = os.path.join(self.base_dir, doc_id)
        if not os.path.isdir(doc_dir):
            return None
        
        # Find document file
        for file in os.listdir(doc_dir):
            if file != "metadata.json":
                return os.path.join(doc_dir, file)
        
        return None
    
    def get_document_metadata(self, doc_id: str) -> Optional[Dict]:
        """
        Get metadata for a stored document.
        
        Args:
            doc_id: Document ID
            
        Returns:
            Document metadata or None if not found
        """
        metadata_path = os.path.join(self.base_dir, doc_id, "metadata.json")
        if not os.path.isfile(metadata_path):
            return None
        
        import json
        with open(metadata_path, 'r') as f:
            return json.load(f)
    
    def list_documents(self) -> List[str]:
        """
        List all stored document IDs.
        
        Returns:
            List of document IDs
        """
        if not os.path.isdir(self.base_dir):
            return []
        
        return [d for d in os.listdir(self.base_dir) 
                if os.path.isdir(os.path.join(self.base_dir, d))]
    
    def delete_document(self, doc_id: str) -> bool:
        """
        Delete a stored document.
        
        Args:
            doc_id: Document ID
            
        Returns:
            True if document was deleted, False otherwise
        """
        doc_dir = os.path.join(self.base_dir, doc_id)
        if not os.path.isdir(doc_dir):
            return False
        
        shutil.rmtree(doc_dir)
        return True
    
    def _generate_document_id(self, file_path: str) -> str:
        """
        Generate a unique document ID based on content hash and timestamp.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document ID
        """
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        # Calculate SHA-256 hash of file content
        with open(file_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()[:16]
        
        return f"{timestamp}_{file_hash}"


class DocumentIngestionManager:
    """Manages the document ingestion process."""
    
    def __init__(self):
        """Initialize the document ingestion manager."""
        self.validator = DocumentValidator()
        self.storage = DocumentStorage()
    
    def ingest_document(self, file_path: str, metadata: Optional[Dict] = None) -> Tuple[bool, str, Optional[str]]:
        """
        Ingest a document: validate and store it.
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata for the document
            
        Returns:
            Tuple of (success, message, document_id)
        """
        # Validate document
        is_valid, message = self.validator.validate_document(file_path)
        if not is_valid:
            return False, message, None
        
        # Store document
        try:
            doc_id = self.storage.store_document(file_path, metadata)
            return True, "Document ingested successfully", doc_id
        except Exception as e:
            return False, f"Error storing document: {str(e)}", None
    
    def batch_ingest_documents(self, file_paths: List[str], metadata_list: Optional[List[Dict]] = None) -> Dict[str, Tuple[bool, str, Optional[str]]]:
        """
        Ingest multiple documents.
        
        Args:
            file_paths: List of paths to document files
            metadata_list: Optional list of metadata for each document
            
        Returns:
            Dictionary mapping file paths to ingestion results
        """
        results = {}
        
        for i, file_path in enumerate(file_paths):
            metadata = metadata_list[i] if metadata_list and i < len(metadata_list) else None
            results[file_path] = self.ingest_document(file_path, metadata)
        
        return results


# Create singleton instances
document_validator = DocumentValidator()
document_storage = DocumentStorage()
document_ingestion = DocumentIngestionManager()
